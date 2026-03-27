#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建优化版的Clash配置教程文档（保留原始图片）
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建新文档
doc = Document()

# ==================== 标题 ====================
title = doc.add_heading('Clash 客户端下载与配置指南', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ==================== 一、下载地址 ====================
doc.add_heading('一、软件下载地址', level=1)

# PC端
p = doc.add_paragraph()
run = p.add_run('PC端下载：')
run.bold = True
doc.add_paragraph('https://74.211.105.94:36322/down/ld8HmZ9QvLls.7z')

# 安卓端
p = doc.add_paragraph()
run = p.add_run('安卓端下载：')
run.bold = True
doc.add_paragraph('https://74.211.105.94:36322/down/exaZJh95IHmp.apk')

# 订阅地址
p = doc.add_paragraph()
run = p.add_run('订阅地址：')
run.bold = True
doc.add_paragraph('http://74.211.105.94:3000/v2rayse_US_2.yaml')

# ==================== 二、PC端配置教程 ====================
doc.add_heading('二、PC端配置教程', level=1)

# ---- 第一步 ----
doc.add_heading('第一步：安装客户端', level=2)
doc.add_paragraph('下载 PC 端安装包并完成安装。')

# ---- 第二步 ----
doc.add_heading('第二步：导入配置', level=2)
for i, step in enumerate(['打开 Clash 软件', '进入「配置」选项', '将订阅地址粘贴到搜索栏中', '点击下载配置', '选择「v2rayse_US.yaml」文件'], 1):
    doc.add_paragraph(f'{i}. {step}')

# 插入 PC 配置图片
doc.add_paragraph()
p = doc.add_paragraph('操作示意图：')
p.runs[0].bold = True
doc.add_picture('/workspace/extracted_images/image1.png', width=Inches(5.5))
doc.add_paragraph()

# ---- 第三步 ----
doc.add_heading('第三步：选择代理节点', level=2)
for i, step in enumerate(['切换到「代理」选项', '选择「全局」或「规则」模式', '选择「vmess+ws+tls」节点', '观察延迟显示（如「673ms」），出现延迟即表示连接成功'], 1):
    doc.add_paragraph(f'{i}. {step}')

# 插入 PC 代理图片
p = doc.add_paragraph('操作示意图：')
p.runs[0].bold = True
doc.add_picture('/workspace/extracted_images/image2.png', width=Inches(5.5))
doc.add_paragraph()

# 模式说明
doc.add_heading('代理模式说明', level=3)
p = doc.add_paragraph()
run = p.add_run('• 全局模式：')
run.bold = True
p.add_run('代理所有流量')

p = doc.add_paragraph()
run = p.add_run('• 规则模式：')
run.bold = True
p.add_run('仅代理国外网站，国内网站直连（推荐）')

# ---- 第四步 ----
doc.add_heading('第四步：启动系统代理', level=2)
for i, step in enumerate(['切换到「常规」选项', '启动「系统代理」'], 1):
    doc.add_paragraph(f'{i}. {step}')

# 插入 PC 系统代理图片
p = doc.add_paragraph('操作示意图：')
p.runs[0].bold = True
doc.add_picture('/workspace/extracted_images/image3.png', width=Inches(5.5))
doc.add_paragraph()

# 可选功能
doc.add_heading('可选功能', level=3)
p = doc.add_paragraph()
run = p.add_run('TUN 模式：')
run.bold = True
p.add_run('开启后可代理所有应用流量，但可能影响网络速度。')

# ==================== 三、重要提示 ====================
doc.add_heading('三、重要提示', level=1)

p = doc.add_paragraph()
run = p.add_run('⚠ 关机前请退出代理')
run.bold = True
p.add_run('，以避免下次开机时无法联网。')

p = doc.add_paragraph()
run = p.add_run('⚠ 如果开机后无法联网')
run.bold = True
p.add_run('，重新开启 Clash 即可恢复正常。')

p = doc.add_paragraph()
run = p.add_run('💡 建议开启「开机启动」')
run.bold = True
p.add_run('功能，可简化操作步骤。')

# ==================== 四、手机端配置教程 ====================
doc.add_heading('四、手机端配置教程（安卓）', level=1)

# ---- 第一步 ----
doc.add_heading('第一步：安装并打开 Clash', level=2)
doc.add_paragraph('在安卓设备上下载并安装 Clash 软件，然后打开应用。')

# 插入手机主界面图片
p = doc.add_paragraph('软件主界面（未启动状态）：')
p.runs[0].bold = True
doc.add_picture('/workspace/extracted_images/image4.jpeg', width=Inches(2.8))
doc.add_paragraph()

# ---- 第二步 ----
doc.add_heading('第二步：导入订阅', level=2)
doc.add_paragraph('点击「配置」进入配置界面，点击「+」图标。')

# 插入创建配置图片
p = doc.add_paragraph('选择导入方式：')
p.runs[0].bold = True
doc.add_picture('/workspace/extracted_images/image5.jpeg', width=Inches(2.8))
doc.add_paragraph()

doc.add_paragraph('选择「从 URL 导入」，填写配置信息：')

# 插入配置编辑图片
p = doc.add_paragraph('按如图填写「名称」和「URL」：')
p.runs[0].bold = True
doc.add_picture('/workspace/extracted_images/image6.jpeg', width=Inches(2.8))
doc.add_paragraph()

# ---- 第三步 ----
doc.add_heading('第三步：选择节点并启动', level=2)
for i, step in enumerate(['返回「配置」界面', '选中「美国线路」（您配置的线路，一定要选中）', '返回首页', '启动 Clash'], 1):
    doc.add_paragraph(f'{i}. {step}')

# 插入配置列表图片
p = doc.add_paragraph('选中「美国线路」：')
p.runs[0].bold = True
doc.add_picture('/workspace/extracted_images/image7.jpeg', width=Inches(2.8))
doc.add_paragraph()

# 插入成功运行图片
p = doc.add_paragraph('启动成功后界面：')
p.runs[0].bold = True
doc.add_picture('/workspace/extracted_images/image8.jpeg', width=Inches(2.8))
doc.add_paragraph()

# ==================== 页脚 ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('=' * 40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('如有疑问，请联系技术支持')

# 保存
doc.save('/workspace/PC端下载地址_优化版.docx')
print('优化版文档已成功生成！')
